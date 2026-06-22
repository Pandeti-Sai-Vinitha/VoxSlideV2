import fitz  # PyMuPDF
import docx2txt
from docx import Document
import os
from typing import List, Tuple, Dict, Any, Optional
import logging
from html.parser import HTMLParser
import html
from config import EXTRACTED_IMAGES_DIR, parse_versioned_basename

logger = logging.getLogger(__name__)

# Import hierarchical context pipeline components
from .document_analyzer import DocumentBlock


class HTMLBlockExtractor(HTMLParser):
    """
    Parse HTML and extract structured blocks with heading hierarchy
    """
    def __init__(self):
        super().__init__()
        self.blocks = []
        self.current_text = []
        self.current_tag = None
        self.page_num = 1
        self.font_size_map = {
            'h1': 32, 'h2': 28, 'h3': 24, 'h4': 20, 'h5': 16, 'h6': 14,
            'p': 12, 'li': 12, 'td': 12
        }
        self.heading_stack = []  # Track current heading level
        self.in_script = False
        self.in_style = False

    def handle_starttag(self, tag, attrs):
        if tag in ('script', 'style'):
            if tag == 'script':
                self.in_script = True
            else:
                self.in_style = True
            return
        
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'td', 'th'):
            # Flush previous text if any
            if self.current_text:
                self._flush_text()
            self.current_tag = tag

    def handle_endtag(self, tag):
        if tag == 'script':
            self.in_script = False
            return
        if tag == 'style':
            self.in_style = False
            return
        
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'td', 'th', 'div'):
            self._flush_text()
            self.current_tag = None

    def handle_data(self, data):
        if self.in_script or self.in_style:
            return
        
        text = data.strip()
        if text and not text.isspace():
            self.current_text.append(html.unescape(text))

    def _flush_text(self):
        if not self.current_text:
            return
        
        text = ' '.join(self.current_text).strip()
        if not text:
            self.current_text = []
            return
        
        tag = self.current_tag or 'p'
        font_size = self.font_size_map.get(tag, 12)
        is_bold = tag.startswith('h')
        is_list = tag == 'li'
        
        # Determine heading level for hierarchy
        heading_level = None
        if tag.startswith('h'):
            heading_level = int(tag[1])
        
        block = DocumentBlock(
            text=text,
            page=self.page_num,
            font_size=float(font_size),
            is_bold=is_bold,
            is_italic=False,
            is_list=is_list,
            heading_level=heading_level if heading_level else None
        )
        self.blocks.append(block)
        self.current_text = []

    def get_blocks(self):
        # Flush any remaining text
        self._flush_text()
        return self.blocks


def extract_text_and_images_from_pdf(pdf_path, images_folder='extracted_images_pdf'):
    """Legacy function - returns flat text (backward compatible)"""
    os.makedirs(images_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    all_text = ""
    image_paths = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        all_text += page.get_text()
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            img_path = os.path.join(images_folder, f"page{page_num+1}_img{img_index+1}.png")
            if pix.n < 5:
                pix.save(img_path)
            else:
                pix1 = fitz.Pixmap(fitz.csRGB, pix)
                pix1.save(img_path)
                pix1 = None
            pix = None
            image_paths.append(img_path)
    return all_text, image_paths

def _resolve_docx_images_folder(docx_path: str, images_folder: str) -> str:
    """Normalize extracted_images_docx destinations for versioned DOCX filenames."""
    from pathlib import Path

    docx_filename = Path(docx_path).stem
    base_name, version_dir = parse_versioned_basename(docx_filename)
    images_root = Path(images_folder)
    root_folder = Path('extracted_images_docx')

    try:
        images_root_resolved = images_root.resolve()
        root_folder_resolved = root_folder.resolve()
    except Exception:
        images_root_resolved = images_root
        root_folder_resolved = root_folder

    if images_root_resolved == root_folder_resolved or images_root == root_folder:
        if version_dir:
            return str(images_root / base_name / version_dir)
        return str(images_root / docx_filename)

    # If the caller already passed extracted_images_docx/<base_name>, use version subfolder when the
    # DOCX filename itself is versioned.
    if images_root_resolved.parent == root_folder_resolved and images_root.name == base_name:
        if version_dir:
            return str(images_root / version_dir)
        return str(images_root)

    # If the caller passed extracted_images_docx/<full_versioned_name>, normalize to nested structure.
    if images_root_resolved.parent == root_folder_resolved and images_root.name == docx_filename:
        if version_dir:
            return str(images_root.parent / base_name / version_dir)
        return str(images_root)

    return str(images_folder)


def extract_text_and_images_from_docx(docx_path, images_folder: str = None):
    """Legacy function - returns flat text (backward compatible)"""
    if images_folder is None:
        images_folder = str(EXTRACTED_IMAGES_DIR)
    organized_images_folder = _resolve_docx_images_folder(docx_path, images_folder)
    os.makedirs(organized_images_folder, exist_ok=True)
    docx2txt.process(docx_path, organized_images_folder)
    image_files = [os.path.join(organized_images_folder, img) for img in os.listdir(organized_images_folder)]
    doc = Document(docx_path)
    all_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return all_text, image_files




def extract_structured_blocks_from_pdf(
    pdf_path: str,
    images_folder: str = 'extracted_images_pdf'
) -> Tuple[List[DocumentBlock], List[str]]:
    """
    Extract structured blocks from PDF with formatting metadata
    
    Returns:
        (blocks: List[DocumentBlock], image_paths: List[str])
    """
    os.makedirs(images_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    blocks = []
    image_paths = []
    block_counter = 0

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Extract text blocks with formatting info
        page_blocks = page.get_text("dict")["blocks"]

        for block_idx, block_data in enumerate(page_blocks):
            if block_data["type"] == 0:  # Text block
                for line in block_data.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        # Extract formatting
                        font_size = span.get("size", 12)
                        font_name = span.get("font", "")

                        # Heuristics for bold (font name contains "bold")
                        is_bold = "bold" in font_name.lower() or font_size > 14

                        # Heuristics for italic
                        is_italic = "italic" in font_name.lower()

                        # Determine if this is a list or special content
                        is_list = text.startswith(("•", "-", "*", "◦")) or any(
                            f"{i}." in text[:3] for i in range(1, 10)
                        )

                        block = DocumentBlock(
                            text=text,
                            page=page_num + 1,
                            font_size=font_size,
                            is_bold=is_bold,
                            is_italic=is_italic,
                            is_list=is_list,
                        )
                        blocks.append(block)
                        block_counter += 1

            elif block_data["type"] == 1:  # Image block
                # Extract image
                try:
                    xref = block_data.get("xref", None)
                    if xref:
                        pix = fitz.Pixmap(doc, xref)
                        img_path = os.path.join(
                            images_folder,
                            f"page{page_num+1}_img{len(image_paths)+1}.png"
                        )
                        if pix.n < 5:
                            pix.save(img_path)
                        else:
                            pix1 = fitz.Pixmap(fitz.csRGB, pix)
                            pix1.save(img_path)
                            pix1 = None
                        pix = None
                        image_paths.append(img_path)

                        # Add caption block
                        block = DocumentBlock(
                            text=f"[Image: {os.path.basename(img_path)}]",
                            page=page_num + 1,
                            is_caption=True,
                        )
                        blocks.append(block)
                except Exception as e:
                    logger.warning(f"Failed to extract image from page {page_num + 1}: {e}")

    logger.info(f"Extracted {len(blocks)} blocks and {len(image_paths)} images from PDF")
    return blocks, image_paths


def extract_structured_blocks_from_docx(
    docx_path: str,
    images_folder: str = None
) -> Tuple[List[DocumentBlock], List[str]]:
    """
    Extract structured blocks from DOCX with formatting metadata
    
    Returns:
        (blocks: List[DocumentBlock], image_paths: List[str])
    """
    if images_folder is None:
        images_folder = str(EXTRACTED_IMAGES_DIR)
    organized_images_folder = _resolve_docx_images_folder(docx_path, images_folder)
    os.makedirs(organized_images_folder, exist_ok=True)
    
    blocks = []
    image_paths = []

    try:
        # Extract images to organized folder
        docx2txt.process(docx_path, organized_images_folder)
        
        # ✅ Rename images to use 1-based numbering (image1.jpg, image2.jpg, etc.)
        extracted_files = sorted([
            img for img in os.listdir(organized_images_folder)
            if os.path.isfile(os.path.join(organized_images_folder, img))
        ])
        
        renamed_image_paths = []
        for idx, original_name in enumerate(extracted_files, start=1):
            original_path = os.path.join(organized_images_folder, original_name)
            ext = os.path.splitext(original_name)[1]  # Get file extension
            new_name = f"image{idx}{ext}"
            new_path = os.path.join(organized_images_folder, new_name)
            
            # Rename if needed
            if original_path != new_path:
                try:
                    os.rename(original_path, new_path)
                    logger.info(f"  Renamed: {original_name} → {new_name}")
                    renamed_image_paths.append(new_path)
                except Exception as e:
                    logger.warning(f"Failed to rename {original_name}: {e}")
                    renamed_image_paths.append(original_path)
            else:
                renamed_image_paths.append(new_path)
        
        image_paths = renamed_image_paths
        logger.info(f"Extracted {len(image_paths)} images to {organized_images_folder}")
    except Exception as e:
        logger.warning(f"Failed to extract images from DOCX: {e}")

    # Extract text with formatting
    doc = Document(docx_path)
    page_num = 1  # DOCX doesn't have explicit pages, approximate

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Extract formatting from paragraph
        font_size = 12  # Default
        is_bold = False
        is_italic = False
        heading_level = None

        if para.runs:
            # Get formatting from first run
            first_run = para.runs[0]
            if first_run.font.size:
                font_size = first_run.font.size.pt
            is_bold = first_run.font.bold if first_run.font.bold is not None else False
            is_italic = first_run.font.italic if first_run.font.italic is not None else False

        # Detect heading levels from paragraph style
        # Word uses styles like "Heading 1", "Heading 2", etc.
        # Only set heading_level, let the analyzer determine actual sizes
        para_style = para.style.name if para.style else ""
        if para_style:
            if "Heading 1" in para_style:
                heading_level = 1
                # Don't force font size, use actual extracted size
            elif "Heading 2" in para_style:
                heading_level = 2
            elif "Heading 3" in para_style:
                heading_level = 3
            elif "Heading 4" in para_style:
                heading_level = 4
        
        # Detect lists
        is_list = para.style.name.startswith("List") or text.startswith(("•", "-", "*", "◦"))

        block = DocumentBlock(
            text=text,
            page=page_num,
            font_size=font_size,
            is_bold=is_bold,
            is_italic=is_italic,
            is_list=is_list,
            heading_level=heading_level,
        )
        blocks.append(block)

    logger.info(f"Extracted {len(blocks)} blocks and {len(image_paths)} images from DOCX")
    return blocks, image_paths


def extract_structured_blocks_from_html(
    html_path: str,
    images_folder: str = 'extracted_images_html'
) -> Tuple[List[DocumentBlock], List[str]]:
    """
    Extract structured blocks from HTML with heading-based hierarchy detection

    Returns:
        (blocks: List[DocumentBlock], image_paths: List[str])
    """
    os.makedirs(images_folder, exist_ok=True)
    image_paths = []

    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
    except UnicodeDecodeError:
        with open(html_path, 'r', encoding='latin-1') as f:
            html_content = f.read()

    # Parse HTML
    parser = HTMLBlockExtractor()
    try:
        parser.feed(html_content)
    except Exception as e:
        logger.warning(f"Error parsing HTML: {e}")

    blocks = parser.get_blocks()

    logger.info(f"Extracted {len(blocks)} blocks from HTML")
    return blocks, image_paths


def extract_with_hierarchy(
    file_path: str,
    file_type: str = "pdf",
    images_folder: Optional[str] = None
) -> Tuple[List[DocumentBlock], Any, List[str]]:
    """
    Extract document and build hierarchy in one step
    
    Args:
        file_path: Path to PDF, DOCX, or HTML
        file_type: "pdf", "docx", or "html"
        images_folder: Folder for extracted images
        
    Returns:
        (blocks: List[DocumentBlock], hierarchy: DocumentSection, image_paths: List[str])
    """
    from .document_analyzer import DocumentAnalyzer

    # Set default images folder
    if images_folder is None:
        if file_type.lower() == "docx":
            images_folder = str(EXTRACTED_IMAGES_DIR)
        else:
            images_folder = f'extracted_images_{file_type}'

    # Extract blocks
    if file_type.lower() == "pdf":
        blocks, image_paths = extract_structured_blocks_from_pdf(file_path, images_folder)
    elif file_type.lower() == "docx":
        blocks, image_paths = extract_structured_blocks_from_docx(file_path, images_folder)
    elif file_type.lower() == "html":
        blocks, image_paths = extract_structured_blocks_from_html(file_path, images_folder)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    # Analyze and build hierarchy
    analyzer = DocumentAnalyzer()
    blocks, hierarchy = analyzer.analyze(blocks)

    logger.info(
        f"Extracted and analyzed {file_path}: "
        f"{len(blocks)} blocks, {hierarchy.section_id if hierarchy else 'no'} hierarchy"
    )

    return blocks, hierarchy, image_paths